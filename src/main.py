"""
    Parsing xml to WORD
"""
import argparse
import shutil
import tempfile
from builtins import enumerate
from contextlib import closing
from lxml import etree
from docxtpl import DocxTemplate
from lxml.etree import iterparse

from xml_mp.element_to_dict import *
from docx import Document
import os
import logging.config

# set up logging
try:
    logging.config.fileConfig("logging_config.ini")
except:
    pass
logger = logging.getLogger('sLogger')


class MpXMlToWORd:
    """
        Преобразователь xml межевого  в ворд
    """
    CNST_FORMAT = 'docx'
    CNST_PATH_TPL = 'template/common/'

    def __init__(self):
        self.name_number = 0
        self.tempfolder = tempfile.mkdtemp()

    def close(self):
        """
            удаление темповой директории
        :return:
        """
        if os.path.exists(self.tempfolder):
            shutil.rmtree(self.tempfolder)

    def fast_iter_element(self, elem: object, func: object, args: object = [], kwargs: object = {}) -> object:
        """
            the node  cleaning
            :param context: context
            :param func: callback - renderToTPL
            :param args: args
            :param kwargs: kwargs
            :return: None
            :rtype: XMLElemet
        """
        func(elem, *args, **kwargs)
        elem.clear()
        while elem.getprevious() is not None:
            if type(elem.getprevious()) == etree._Element:
                if elem.getparent() is not None:
                    del elem.getparent()[0]
            else:
                break

    def run_render_tpl(self, elem, xml_class_name, is_clean, pos_node):
        """
        Запуск парсинга определенного блока xml
        :param elem: node
        :param xml_class_name: class -> retun dict
        :param is_clean: очищать  узел или там еще что то нужно
        :param pos_node: просто порядковый номер позици узла
        :return: docx
        """
        if is_clean:
            self.fast_iter_element(elem, self.render_tpl, args=(xml_class_name,
                                                                self.CNST_PATH_TPL + BINDER_FILE[elem.tag]['tpl'],
                                                                BINDER_FILE[elem.tag]['pos_doc'] + str(pos_node)))
        else:
            self.render_tpl(elem, xml_class_name, self.CNST_PATH_TPL + BINDER_FILE[elem.tag]['tpl'],
                                                                BINDER_FILE[elem.tag]['pos_doc'] + str(pos_node))

    def __context_parser(self, context):
        """
            Парсим node
        :param context:
        :return:None
        """
        i = 0
        try:
            for event, elem in context:
                i += 1
                if elem.tag in cnfg.LIST_BLOCK_NODE and event == 'end':
                    if elem.tag == 'SubParcels' and event == 'end' and elem.getparent().tag != 'Package':
                        continue
                    self.run_render_tpl(elem, BINDER_FILE[elem.tag]['class'], BINDER_FILE[elem.tag]['clear'], i)
        except Exception as e:
            logger.error(f"""{e} ->{elem} """)
        finally:
            del context

    def __xml_block_to_docx(self, path):
        """
            Формирование списка док. файлов  по блокам xml
        :param path: путь до xml файла
        """
        # get an iterable
        context = iterparse(path, events=("start", "end"))
        context = iter(context)
        self.__context_parser(context)
        del context

    def render_tpl(self,node, XMLClass, path_tpl, name_result):
        """
            Рендер шаблона
        :param node:  узел- noda
        :param XMLClass: класс отвечающий за парсинг данной ноды в dict (to_dict)
        :param path_tpl: путь до template
        :return: word файл  с наименованием =  [Number - позиция word- элемента в файле]+ [Number - позиция node].docx
        """
        try:
            if len(node) > 0 or node.text:
                tpl = DocxTemplate(path_tpl)
                instance = XMLClass(node)
                tpl.render(instance.to_dict())
                file_res = '.'.join([name_result, self.CNST_FORMAT])
                tpl.save(os.path.join(self.tempfolder, file_res))
                logger.info(f"""Parsing {node}  done -> result {name_result}""")
        except Exception as e:
            logger.error(f"""Error parsing {node} : {e}""")

    def __element_body_docx(self, path):
        """
        :param path: получить блок ворд -файла
        :return: element docx
        """
        doc = Document(path)
        for element in doc.element.body:
            yield element

    def combine_word_documents(self, result_path_file):
        """
        Собираем все файлы в единый документ
        :param input_files: iterable список файлов
        :return: result braid Docx
        """
        files = sorted(os.listdir(self.tempfolder))
        _dcx = filter(lambda x: x.endswith('.' + self.CNST_FORMAT), files)
        _dcx = map(lambda x: os.path.join(self.tempfolder, x), _dcx)

        merged_document = Document()
        for filnr, file in enumerate(_dcx):
            _ = os.path.join(file)
            if filnr == 0:
                merged_document = Document(_)
            else:
                for element in self.__element_body_docx(_):
                    merged_document.element.body.append(element)
                # if filnr < len(files) - 1:
                #     merged_document.element.body.text.page_break_before()

            # if filnr == 0:
            #     merged_document = Document(_)
            #     merged_document.add_page_break()
            # else:
            #     for el in self.__element_body_docx(_):
            #         merged_document.element.body.append(el)
            #     merged_document.add_page_break()

        merged_document.save(result_path_file)

    def run(self, path_file, result_file):
        """
            run convert xml to  word
        :param path_file:  sourse file xml
        :param result_file:  path file result
        :return:
        """
        generat.__xml_block_to_docx(path_file)
        logger.info('START COMBINE WORDS')
        generat.combine_word_documents(result_file)
        logger.info('END')


if __name__ == '__main__':
    """
        поехали 
    """
    logger.info('START PARSING')
    parser = argparse.ArgumentParser()

    parser.add_argument('-i', '--input', help='Путь к xml-файлу', type=str)
    parser.add_argument('-o', '--output', help='Файл для получения реультата', type=str)
    args = parser.parse_args()
    if not args.input:
        parser.print_help()
    else:
        xml_file = os.path.normpath(args.input)
        res_file = os.path.normpath(args.output)
        try:
            with closing(MpXMlToWORd()) as generat:
                generat.run(xml_file, res_file)
        except Exception as e:
            logger.error(f"""Error parsing file {xml_file}  -> {e}""")
        else:
            logger.info(f"""The file path = {xml_file} -> parsing done! """)