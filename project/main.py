import shutil
import tempfile
from builtins import enumerate
from contextlib import closing
from lxml import etree
from docxtpl import DocxTemplate
from lxml.etree import iterparse

from core.xml_to_dict import  *
from docx import Document
import os

import logging
from logging.config import fileConfig
fileConfig('loggers/logging_config.ini')
logger = logging.getLogger()


class XMLIter__context_parser:

    def __init__(self, pathxml):
        self.pathxml = pathxml
        self.context = None

    context = property()

    @context.getter
    def context(self):
        self.context = iter(etree.iterparse(self.pathxml, events=('start', 'end',)))
        for event, element in self.context:
            yield event, element

    def __del__(self):
        if self.context:
            for event, element in self._context:
                if event == 'end':
                    element.clear()
                    while element.getprevious() is not None:
                        if type(element.getprevious()) == etree._Element:
                            if element.getparent() is not None:
                                del element.getparent()[0]
                        else:
                            break
        del self.context

    def getcontext(self):
        if not self._context:
            self._context = iter(etree.iterparse(self.xml, events=('start', 'end',)))
        for event, element in self._context:
            self._current = element
            yield event, element

    def delcontext(self):
        del self._context


class MpXMlToWORd:

    def __init__(self):
        self.name_number = 0
        self.tempfolder = tempfile.mkdtemp()

    def close(self):
        if os.path.exists(self.tempfolder):
            shutil.rmtree(self.tempfolder)

    def fast_iter(self,context, func, args=[], kwargs={}):
        # http://www.ibm.com/developerworks/xml/library/x-hiperfparse/
        for event, elem in context:
            # Also eliminate now-empty references from the root node to elem
            for ancestor in elem.xpath('ancestor-or-self::*'):
                print('Checking ancestor: {a}'.format(a=ancestor.tag))
                while ancestor.getprevious() is not None:
                    print(
                        'Deleting {p}'.format(p=(ancestor.getparent()[0]).tag))
                    del ancestor.getparent()[0]
            print('Processing {e}'.format(e=etree.tostring(elem)))
            func(elem, *args, **kwargs)
            # It safe to call clear() here because no descendants will be
            # accessed
            print('Clearing {e}'.format(e=etree.tostring(elem)))
            elem.clear()

        del context

    def __fast_iter_element(self, elem: object, func: object, args: object = [], kwargs: object = {}) -> object:
        """

        :rtype: XMLElemet
        """
        func(elem, *args, **kwargs)
        elem.clear()
        while elem.getprevious() is not None:
            if type(elem.getprevious()) == etree._Element:
                if elem.getparent() is not None:
                    del elem.getparent()[0]
            else:
                break

    def __context_parser(self, context):
        i = 0
        for event, elem in context:
            i += 1

            if elem.tag == 'GeneralCadastralWorks' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlTitleDict, 'template/common/title.docx','1.' + str(i)))
            if elem.tag == 'InputData' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlInputDataDict, 'template/common/inputdata.docx', '2.' + str(i)))
            if elem.tag == 'Survey' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlSurveyDict, 'template/common/survey.docx','3.' + str(i)))
            if elem.tag == 'NewParcel' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlNewParcel, 'template/common/newparcel.docx', '4.' + str(i)))
            if elem.tag == 'ExistParcel' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlExistParcel, 'template/common/existparcel.docx', '4.' + str(i)))

            if elem.tag == 'SubParcels' and event == 'end' and elem.getparent().tag != 'InputData':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlSubParcels, 'template/common/subparcels.docx', '6.' + str(i)))

            if elem.tag == 'ChangeParcel' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlChangeParcel, 'template/common/changeparcel.docx', '5.' + str(i)))

            if elem.tag == 'SpecifyRelatedParcel' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlExistParcel, 'template/common/existparcel.docx','6.' + str(i)))
            if elem.tag == 'FormParcels' and event == 'start':
                logging.info(elem.tag)
                self.render_tpl(elem,XmlNewParcelProviding, 'template/common/providing.docx','7.' + str(i))

            if elem.tag == 'Conclusion' and event == 'end':
                logging.info(elem.tag)
                self.__fast_iter_element(elem, self.render_tpl,
                                       args=(XmlConclusion, 'template/common/conclusion.docx','8.' + str(i)))
            _elin =elem.xpath('InputData/Documents')
            if _elin:
                logging.info(f"Documents est")
            else:
                logging.info(" nett")

        del context

    def __xml_block_to_docx(self, path):
        """
            Формирование списка док. файлов  по блокам xml
        :param path: путь до файла
        :return:
        """
        # get an iterable
        context = iterparse(path, events=("start", "end"))
        context = iter(context)
        self.__context_parser(context)
        del context

    def render_tpl(self,node, XMLClass, path_tpl, name_result):
        """
            Рендер шаблона
        :param node:  узел- noda
        :param XMLClass: класс отвечающий за парсинг данной ноды в dict (to_dict)
        :param path_tpl: путь до template
        :return: file docx
        """
        if len(node) > 0 or node.text:
            tpl = DocxTemplate(path_tpl)
            instance = XMLClass(node)
            tpl.render(instance.to_dict())
            file_res = '.'.join([name_result, 'docx'])
            tpl.save(os.path.join(self.tempfolder, file_res))

    def __get_element_body_docx(self, path):
        """
        :param path: блок ворд -файла
        :return:
        """
        doc = Document(path)
        for element in doc.element.body:
            yield element

    def combine_word_documents(self, result_path_file):
        """
        :param input_files: iterable список файлов
        :return: Docx
        """
        files = os.listdir(self.tempfolder)
        _dcx = filter(lambda x: x.endswith('.docx'), files)
        _dcx = map(lambda x: os.path.join(self.tempfolder, x), _dcx)

        merged_document = Document()
        for filnr, file in enumerate(_dcx):
            # if 'offerte_template' in file:
            #     file = os.path.join(file)
            _ = os.path.join(file)
            if filnr == 0:
                merged_document = Document(_)
            else:
                for element in self.__get_element_body_docx(_):
                    merged_document.element.body.append(element)
                # if filnr < len(files) - 1:
                #     merged_document.element.body.text.page_break_before()

            # if filnr == 0:
            #     merged_document = Document(_)
            #     merged_document.add_page_break()
            # else:
            #     for el in self.__get_element_body_docx(_):
            #         merged_document.element.body.append(el)
            #     merged_document.add_page_break()

        merged_document.save(result_path_file)

    def run(self, path_file, result_file):
        generat.__xml_block_to_docx(path_file)
        logger.info('START COMBINE WORDS')
        generat.combine_word_documents(result_file)
        logger.info('END')


if __name__ == '__main__':
    logger.info('START PARSING')
    try:
        with closing(MpXMlToWORd()) as generat:
            generat.run('../TEST/3/3.xml', '../TEST/3/3.docx')
    except Exception as e:
        logger.error(f"""Error parsing file {e}""")
    else:
        logger.info('The file parsing done!')
